# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


def build_document():
    doc = Document()

    # Márgenes
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Tipografía base
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Título
    title = doc.add_heading('Informe Técnico – Aplicación “Calendario Semanal (Flutter Web)”', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Sección 1
    doc.add_heading('1. Introducción', level=1)
    doc.add_paragraph(
        'Este documento describe el funcionamiento y la arquitectura de una aplicación web desarrollada en Flutter '
        'que permite gestionar eventos en una vista semanal. El objetivo es presentar, de forma clara y técnica, '
        'cómo está construida la solución, qué problemas resuelve y qué decisiones de diseño se tomaron.'
    )
    doc.add_paragraph(
        'Público objetivo: docentes, evaluadores y compañeros de Tecsup interesados en comprender el diseño, '
        'la lógica y el flujo de la app.'
    )

    # Sección 2
    doc.add_heading('2. Resumen funcional', level=1)
    bullets = [
        'Vista semanal (lunes a domingo) con navegación entre semanas.',
        'CRUD de eventos: crear, ver detalle, editar y eliminar.',
        'Repetición de eventos: solo ese día, semanal o días específicos (p. ej., lunes y sábado).',
        'Chip de importancia (Baja, Media, Alta) con colores distintivos.',
        'Animación de deslizamiento fluida y direccional al cambiar de semana.',
        'Soporte para eventos de “todo el día” o con horario definido.',
        'Interfaz en español con diseño Material 3 y tipografía Poppins.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Sección 3
    doc.add_heading('3. Tecnologías y arquitectura', level=1)
    bullets = [
        'Framework: Flutter (Web) y lenguaje Dart.',
        'Gestión de estado: Provider con ChangeNotifier.',
        'UI/UX: Material Design 3, Google Fonts (Poppins).',
        'Utilidades: uuid para IDs únicos, intl opcional para formateos.',
        'Estructura en capas: models, providers, data, screens/widgets, theme.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Sección 4
    doc.add_heading('4. Estructura del proyecto (resumen)', level=1)
    bullets = [
        'lib/main.dart: punto de entrada; inyecta el Provider y el tema.',
        'lib/models/event_model.dart: modelo CalendarEvent y enums auxiliares.',
        'lib/providers/event_provider.dart: estado global, CRUD y navegación semanal.',
        'lib/data/sample_events.dart: datos simulados (3 semanas).',
        'lib/screens/weekly_calendar_screen.dart: pantalla principal con animación.',
        'lib/widgets/: tarjetas de evento, formularios y diálogos de detalle.',
        'lib/theme/app_theme.dart: Material 3, colores y tipografía.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    doc.add_paragraph('(Insertar imagen aquí: mapa de carpetas o diagrama general)').italic = True

    # Sección 5
    doc.add_heading('5. Modelo de datos', level=1)
    doc.add_paragraph('Entidad principal: CalendarEvent')
    bullets = [
        'id: String (uuid v4).',
        'title: String.',
        'description: String.',
        'location: String? (opcional).',
        'date: DateTime (fecha base de referencia).',
        'startTime / endTime: TimeOfDay? (nulos si es “todo el día”).',
        'color: Color (personalización visual).',
        'category: EventCategory {exams, work, task, birthday, other}.',
        'repeatType: RepeatType {single, weekly, custom}.',
        'repeatDays: List<int> (1=Lun … 7=Dom; usado cuando repeatType=custom).',
        'importance: EventImportance {low, medium, high}.',
        'isAllDay: getter que deriva si no hay horas definidas.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    doc.add_paragraph(
        'Este diseño permite inmutabilidad con copyWith, flexibilidad para representar distintos patrones de repetición '
        'y visualización consistente en la UI (categoría e importancia definen íconos/colores).'
    )

    # Sección 6
    doc.add_heading('6. Lógica de negocio clave', level=1)
    doc.add_paragraph('Filtro de eventos por fecha (getEventsForDate):')
    bullets = [
        'Coincidencia exacta: event.date == día seleccionado.',
        'Repetición semanal: coincide si weekday(event) == weekday(día objetivo) y el día objetivo es el mismo o posterior a event.date.',
        'Repetición por días específicos (custom): coincide si repeatDays contiene weekday(día objetivo) y el día objetivo es el mismo o posterior a event.date.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    doc.add_paragraph('Ordenamiento por día: todo el día primero y luego por hora de inicio ascendente.')

    # Sección 7
    doc.add_heading('7. Gestión de estado (Provider)', level=1)
    bullets = [
        'Mantiene: _events (lista maestra), _currentWeekStart (lunes visible), _slideDirection (dirección de animación).',
        'Operaciones: goToPreviousWeek, goToNextWeek, goToCurrentWeek.',
        'Servicios: getDaysOfCurrentWeek, getEventsForDate, add/update/delete.',
        'notifyListeners() para refrescar la UI tras cambios.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')
    doc.add_paragraph('(Insertar imagen aquí: diagrama simple del flujo Provider → Widgets)').italic = True

    # Sección 8
    doc.add_heading('8. Interfaz de usuario y flujos', level=1)
    bullets = [
        'Pantalla principal: AppBar con navegación y botón “Hoy”, encabezado de días (Lun–Dom) y columnas con EventCard.',
        'Formulario (EventFormDialog): título, descripción, ubicación, fecha, horario/todo el día, color, categoría e importancia.',
        'Repetición: “solo ese día”, “semanal” o “días específicos” (chips Lun..Dom).',
        'Diálogo de detalle: muestra toda la información y acciones de Editar/Eliminar (con confirmación).',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Sección 9
    doc.add_heading('9. Animaciones de cambio de semana', level=1)
    bullets = [
        'AnimationController (400 ms, curva easeOutCubic).',
        'SlideTransition con desplazamiento horizontal (±0.3 → 0) y FadeTransition (0 → 1).',
        'ClipRect para evitar desbordes visuales.',
        'El disparo ocurre al detectar cambio de currentWeekStart; la dirección proviene de slideDirection.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Sección 10
    doc.add_heading('10. Datos de ejemplo', level=1)
    doc.add_paragraph(
        'Se incluyen tres semanas (anterior, actual y siguiente) con eventos variados en español, '
        'asignados a distintas categorías e importancias, y con patrones de repetición heterogéneos. '
        'Esto permite validar visual y funcionalmente la aplicación sin necesidad de backend.'
    )

    # Sección 11
    doc.add_heading('11. Pruebas', level=1)
    doc.add_paragraph(
        'Se incluye una prueba de humo (widget test) que verifica el render de la app mediante la presencia '
        'del botón “Nuevo”. Sirve como base para extender con pruebas de interacción y de lógica.'
    )

    # Sección 12
    doc.add_heading('12. Ejecución y despliegue', level=1)
    doc.add_paragraph('Desarrollo:')
    for b in ['flutter pub get', 'flutter run -d web-server --web-port=8080']:
        doc.add_paragraph(b, style='List Bullet')
    doc.add_paragraph('Producción:')
    for b in ['flutter build web (genera build/web listo para hosting estático)']:
        doc.add_paragraph(b, style='List Bullet')

    # Sección 13
    doc.add_heading('13. Buenas prácticas aplicadas', level=1)
    bullets = [
        'Separación por capas (modelos, estado, UI, tema).',
        'Inmutabilidad (copyWith) y patrones claros de datos.',
        'Consistencia visual (Material 3, Poppins, esquema de color).',
        'Componentización de la UI (EventCard, diálogos).',
        'Legibilidad del código y nombres autoexplicativos.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Sección 14
    doc.add_heading('14. Limitaciones y trabajo futuro', level=1)
    bullets = [
        'Persistencia: actualmente en memoria (no hay SQLite/Firestore/REST).',
        'Colisiones/overlaps: no hay layout inteligente para eventos solapados.',
        'Arrastrar y soltar (drag & drop): no implementado.',
        'Notificaciones recordatorias: fuera de alcance.',
        'Soporte multi-idioma y accesibilidad: posible mejora.',
        'Más pruebas unitarias y de integración.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Sección 15
    doc.add_heading('15. Conclusiones', level=1)
    doc.add_paragraph(
        'La aplicación cumple con los objetivos de un calendario semanal con CRUD, repetición flexible y una '
        'experiencia cuidada, incluyendo animaciones. La arquitectura con Provider permite mantener la lógica '
        'desacoplada de la UI, facilitando futuras extensiones como persistencia o integración con servicios externos.'
    )

    # Sección 16
    doc.add_heading('16. Glosario breve', level=1)
    bullets = [
        'Provider/ChangeNotifier: patrón de gestión de estado reactivo en Flutter.',
        'Widget: bloque fundamental de UI en Flutter.',
        'Enum: tipo enumerado para modelar categorías, importancia y patrones de repetición.',
        'Material 3: guía de diseño de Google para interfaces modernas.',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    # Guardar en carpeta docs
    os.makedirs('docs', exist_ok=True)
    out_path = os.path.join('docs', 'AMM Lab06 ANavarro 5C24A.docx')
    doc.save(out_path)
    return out_path


if __name__ == '__main__':
    path = build_document()
    print(f'Documento generado: {path}')
